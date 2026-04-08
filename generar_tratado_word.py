import os
import sys

# Auto-instalación de python-docx si no existe
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    print("Instalando librerias necesarias...")
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# CONFIGURACIÓN DE PÁGINA - ESTILO DIPLOMÁTICO
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def add_spacer(n=1):
    for _ in range(n): doc.add_paragraph()

def set_style(run, size=11, font='Georgia', color=RGBColor(26, 26, 26), bold=False, italic=False):
    run.font.size = Pt(size)
    run.font.name = font
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic

def add_section_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_style(run, size=13, font='Arial', color=RGBColor(0, 35, 149), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_body(text, bold_start=False):
    p = doc.add_paragraph(style='List Bullet')
    if bold_start and ":" in text:
        parts = text.split(":", 1)
        set_style(p.add_run(parts[0] + ":"), bold=True)
        set_style(p.add_run(parts[1]))
    else:
        set_style(p.add_run(text))

# ==========================================
# PÁGINA 1: PORTADA SOBERANA
# ==========================================
add_spacer(2)
p_label = doc.add_paragraph()
p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_style(p_label.add_run("DOCUMENTO SOBERANO"), size=10, font='Arial', color=RGBColor(133, 113, 77), bold=True)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("TRATADO DE\nCOOPERACIÓN\nESTRATÉGICA\n2026")
set_style(run_title, size=36, font='Arial Black', color=RGBColor(0, 35, 149), bold=True)

add_spacer(3)
p_entities = doc.add_paragraph()
p_entities.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_style(p_entities.add_run("ALLIANCE FRANÇAISE DE SAN CRISTÓBAL DE LAS CASAS\n"), size=14, font='Arial', bold=True)
set_style(p_entities.add_run("&\n"), size=28, font='Georgia', color=RGBColor(179, 0, 0), bold=True) # Rojo AF
set_style(p_entities.add_run("INSTITUTO DE ARTES CULINARIAS PARMENTIER"), size=14, font='Arial', bold=True)

add_spacer(4)
p_loc = doc.add_paragraph()
p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_style(p_loc.add_run("PÁGINA I | UNIDAD DE ALTA DIRECCIÓN GENERAL"), size=9, color=RGBColor(133, 113, 77), bold=True)
doc.add_page_break()

# ==========================================
# PÁGINA 2: EXORDIO Y DECLARACIONES
# ==========================================
add_section_header("I. EXORDIO INSTITUCIONAL Y ANTECEDENTES")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_style(p.add_run(
    "El presente Tratado es la materialización de un esfuerzo ininterrumpido. "
    "Ambas partes declaran que desde el año 2023 sostienen una colaboración "
    "académica y cultural que ha dado frutos tangibles, tales como la integración "
    "del francés en el plan formativo del Instituto, la creación de proyectos conjuntos "
    "como el huerto urbano con fines pedagógicos, y la organización de talleres "
    "culinarios y agroecológicos en lengua francesa. Este documento 2026 constituye "
    "una renovación integral y estratégica de dichos acuerdos, elevando el rigor "
    "académico hacia la movilidad internacional."
))

add_section_header("II. DECLARACIONES DE LEY")
p = doc.add_paragraph()
set_style(p.add_run("2.1 DECLARA «LA ALIANZA»:"), bold=True)
add_body("Que la Alianza Franco Mexicana de San Cristóbal, A.C. está legalmente "
         "constituida como Asociación Civil. RFC: AFM110210ER4.")
add_body("Que su domicilio fiscal se ubica en Av. La Almolonga N° 80, Barrio Santa Lucía, "
         "C.P. 29250, San Cristóbal de las Casas, Chiapas.")
add_body("Que sus datos de contacto oficiales son el correo: afsancris@gmail.com "
         "y el teléfono institucional: 967 172 1870.")
add_body("Que se encuentra representada por su Presidenta del Consejo Directivo, la "
         "Lic. Diana Verania Zebadúa Coello, y su Director General, el Lic. Miguel David Tillero Álvarez.")

p = doc.add_paragraph()
set_style(p.add_run("2.2 DECLARA «EL INSTITUTO»:"), bold=True)
add_body("Que el Instituto de Artes Culinarias Parmentier de San Cristóbal, S.C. es una "
         "institución de educación técnica superior. RFC: IAC1006059Q5.")
add_body("Que se encuentra debidamente representado por su Directora, la Lic. Cynthia "
         "Guadalupe Martínez Orantes.")
doc.add_page_break()

# ==========================================
# PÁGINA 3: ALCANCE Y COMPROMISOS
# ==========================================
add_section_header("III. ALCANCE GENERAL Y OBJETO")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_style(p.add_run("PRIMERA. OBJETO: Regular la prestación de servicios académicos por "
                    "parte de «LA ALIANZA» a «EL INSTITUTO», abarcando la enseñanza del idioma "
                    "francés técnico gastronómico y agroecológico (Nivel A1.1)."), bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_style(p.add_run("SEGUNDA. ALCANCE Y HORARIOS: «EL INSTITUTO» contrata un programa "
                    "intensivo de 30 horas lectivas. Las clases se impartirán los días viernes, "
                    "de 09:00 a 11:00 AM, del 16 de enero al 08 de mayo de 2026."), bold=True)

add_section_header("IV. COMPROMISOS INSTITUCIONALES (AFSC)")
add_body("Gestión Docente: Designación del Maestro José Antonio Miranda, garantizando excelencia técnica.", bold_start=True)
add_body("Reporte y Seguimiento: Entrega mensual de evaluaciones y asistencia a la Dirección.", bold_start=True)
add_body("Programa de Movilidad: Asesoría activa para la postulación a inmersión en Francia.", bold_start=True)
doc.add_page_break()

# ==========================================
# PÁGINA 4: PLANIFICACIÓN Y HUERTO
# ==========================================
add_section_header("V. PLANIFICACIÓN PEDAGÓGICA Y MÓDULOS")
p = doc.add_paragraph()
set_style(p.add_run("El programa de francés técnico (30 horas) estará estructurado en 4 módulos progresivos:"))
add_body("MÓDULO 1: VOCABULARIO TÉCNICO Y FUNDAMENTOS (Semanas 1-4). Léxico de cocina, huerto y jerarquía de brigada.", bold_start=True)
add_body("MÓDULO 2: GESTIÓN DE INGREDIENTES Y RECETAS (Semanas 5-8). Redacción y traducción de fichas técnicas.", bold_start=True)
add_body("MÓDULO 3 Y 4: SERVICIO DE SALA E INMERSIÓN (Semanas 9-15). Protocolo internacional y preparación cultural Vichy.", bold_start=True)

add_section_header("VI. PROGRAMA PRÁCTICO (EL HUERTO URBANO AFSC)")
p = doc.add_paragraph()
set_style(p.add_run("Tres visitas obligatorias a la AFSC para fusionar lingüística con praxis agroecológica:"))
add_body("VISITA INICIAL (Febrero 2026): Reconocimiento y siembra inicial en francés.", bold_start=True)
add_body("VISITA INTERMEDIA (Marzo 2026): Mantenimiento y planificación de menú orgánico.", bold_start=True)
add_body("VISITA FINAL (08 Mayo 2026): Cosecha y Taller Gastronómico 100% en francés.", bold_start=True)
doc.add_page_break()

# ==========================================
# PÁGINA 5: FINANZAS Y MOVILIDAD
# ==========================================
add_section_header("VII. CONDICIONES ECONÓMICAS Y FINANCIERAS")
p = doc.add_paragraph()
set_style(p.add_run("Bajo el seguimiento de la C.P. Berenice Díaz, se establece el esquema financiero:"))
add_body("INVERSIÓN TOTAL: $9,300.00 MXN ($310.00 MXN por hora).", bold_start=True)
add_body("PRIMER PAGO (30%): $2,790.00 MXN (Antes del 30 Ene 2026).", bold_start=True)
add_body("SEGUNDO PAGO (40%): $3,720.00 MXN (Antes del 13 Mar 2026).", bold_start=True)
add_body("TERCER PAGO (30%): $2,790.00 MXN (Antes del 08 May 2026).", bold_start=True)

add_section_header("VIII. PROYECCIÓN INTERNACIONAL (CAVILAM-VICHY)")
p = doc.add_paragraph()
set_style(p.add_run("SEXTA. MOVILIDAD: El presente Tratado es la piedra angular para la "
                    "internacionalización estudiantil hacia CAVILAM – Alliance Française en Vichy, "
                    "Francia (Septiembre 2026). Exclusiones: Vuelos y viáticos son independientes."), bold=True)

add_section_header("IX. MARCO LEGAL Y PRIVACIDAD")
add_body("PROPIEDAD INTELECTUAL: El uso de logotipos requiere autorización del Director General de la AFSC.")
add_body("PROTECCIÓN DE DATOS: Estricto apego a la LFPDPPP.")
doc.add_page_break()

# ==========================================
# PÁGINA 6: ANEXO I (CRONOGRAMA)
# ==========================================
add_section_header("ANEXO I: CRONOGRAMA EXACTO Y TALLER FINAL")
p = doc.add_paragraph()
set_style(p.add_run("Las 15 sesiones se ejecutarán en los siguientes viernes de 2026:"))
add_body("ENERO: 16 (Inicio), 23 y 30.", bold_start=True)
add_body("FEBRERO: 06, 13 y 27 (20 Suspensión).", bold_start=True)
add_body("MARZO: 06, 13, 20 y 27.", bold_start=True)
add_body("ABRIL: 10, 17 y 24 (03 Feriado).", bold_start=True)
add_body("MAYO: 08 (Taller Final de Evaluación).", bold_start=True)

p = doc.add_paragraph()
set_style(p.add_run("\nLOGÍSTICA DEL TALLER CULINARIO (08 DE MAYO 2026)"), bold=True)
add_body("09:00 - 09:30 AM: Cosecha y Mise en place (Francés).")
add_body("09:30 - 10:30 AM: Elaboración gastronómica con léxico técnico.")
add_body("10:30 - 11:00 AM: Degustación y entrega de constancias.")
doc.add_page_break()

# ==========================================
# PÁGINA 7: ANEXO II (PADRÓN Y FIRMAS)
# ==========================================
add_section_header("ANEXO II: PADRÓN OFICIAL A1.1 Y RATIFICACIÓN")
p = doc.add_paragraph()
set_style(p.add_run("Registro oficial de los 12 estudiantes matriculados:"))

alumnos = [
    "GÓMEZ LÓPEZ, ANDREA NAHIELI", "HERNÁNDEZ GONZÁLEZ, JAQUELINE ESMERALDA",
    "LÓPEZ MUÑOZ, MARIO DE JESÚS", "LÓPEZ OCHOA, ABIGAIL",
    "MARTÍNEZ AMBROCIO, JEHOSUA ISAAC", "MÉNDEZ SÁNCHEZ, MARTIN EUSEBIO",
    "PEÑATE GUTIÉRREZ, JEHU HANIEL", "ROMERO LORENZO, SERGIO RODRIGO",
    "RUIZ AGUILAR, ANDREA ITZEL", "RUIZ GONZÁLEZ, ERICK",
    "SANTIZ PÉREZ, HENRY ALBERTO", "VILLALOBOS ESPINOZA, VALERIA"
]

for idx, a in enumerate(alumnos, 1):
    doc.add_paragraph(f"{idx:02d}. {a}", style='List Number')

add_spacer(2)
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_style(p_fin.add_run("Leído el presente instrumento, las partes lo firman en absoluta conformidad "
                        "por duplicado en San Cristóbal de Las Casas, Chiapas, a los 16 días del mes de enero de 2026."), italic=True)

add_spacer(2)
table_sig = doc.add_table(rows=1, cols=2)
table_sig.rows[0].cells[0].text = "_________________________________\nLic. Miguel David Tillero Álvarez\nDirector General AFSC"
table_sig.rows[0].cells[1].text = "_________________________________\nLic. Cynthia G. Martínez Orantes\nDirectora Instituto Parmentier"

for cell in table_sig.rows[0].cells:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_style(run, size=10, bold=True)

add_spacer(2)
p_testigo = doc.add_paragraph()
p_testigo.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_style(p_testigo.add_run("_________________________________\nLic. Diana Verania Zebadúa Coello\nPresidenta del Consejo Directivo AFSC\nTestigo de Honor"), size=10, bold=True)

filename = "Tratado_Magna_AFSC_Parmentier_2026_Final.docx"
doc.save(filename)
os.startfile(filename)

print(f"\n[EXITO] Documento {filename} generado y abierto correctamente en Word.\n")
